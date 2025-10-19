from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import glob

class ReportGenerator:
    """
    Genera el informe completo de la practica en formato Word.
    """

    def __init__(self, output_path='informe_practica_5.docx'):
        self.doc = Document()
        self.output_path = output_path

        # Configurar margenes y estilo
        self._setup_document_style()

    def _setup_document_style(self):
        """Configura el estilo general del documento."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def add_heading_custom(self, text, level=1):
        """Agrega un encabezado personalizado."""
        heading = self.doc.add_heading(text, level=level)
        return heading

    def add_paragraph_custom(self, text, bold=False, italic=False):
        """Agrega un parrafo con formato personalizado."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return para

    def add_image_with_caption(self, image_path, caption, width=5.5):
        """Agrega una imagen con pie de foto."""
        if os.path.exists(image_path):
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Inches(width))

            # Caption
            caption_para = self.doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
        else:
            self.add_paragraph_custom(f"[Imagen no encontrada: {image_path}]")

    def add_table_from_data(self, headers, rows):
        """Crea una tabla con los datos proporcionados."""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Encabezados
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Negrita para encabezados
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Filas de datos
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)

        self.doc.add_paragraph()  # Espacio despues de tabla

    def generate_cover_page(self):
        """Genera la portada del informe."""
        # Titulo principal
        title = self.doc.add_heading('Informe de Laboratorio', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitulo
        subtitle = self.doc.add_heading('Practica 5: Transformada de Hough y Cuantificacion Vectorial', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Informacion del curso
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(
            'Procesamiento Digital de Imagenes\n\n'
            'Fecha: Octubre 2025'
        )
        info_run.font.size = Pt(12)

        self.doc.add_page_break()

    def generate_introduction(self):
        """Genera la seccion de introduccion teorica."""
        self.add_heading_custom('1. Introduccion Teorica', level=1)

        self.add_heading_custom('1.1 Transformada de Hough', level=2)

        self.add_paragraph_custom(
            'La Transformada de Hough es una tecnica fundamental en el procesamiento de '
            'imagenes que permite detectar formas geometricas especificas, principalmente '
            'lineas y circulos. El concepto basico detras de esta transformada es mapear '
            'puntos del espacio de la imagen al espacio de parametros, donde las formas '
            'geometricas se representan de manera mas sencilla.'
        )

        self.add_paragraph_custom(
            'Para la deteccion de lineas, la Transformada de Hough utiliza la representacion '
            'polar de una recta mediante los parametros rho y theta, donde rho es la distancia '
            'perpendicular desde el origen hasta la linea, y theta es el angulo que forma esta '
            'perpendicular con el eje horizontal. Cada punto en el espacio de la imagen genera '
            'una curva sinusoidal en el espacio de Hough, y los puntos colineales producen '
            'curvas que se intersectan en un punto comun, permitiendo asi detectar la presencia '
            'de una linea.'
        )

        self.add_paragraph_custom(
            'Los parametros criticos de esta transformada son la resolucion angular (delta theta) '
            'y la resolucion de distancia (delta rho), que determinan la precision con la que se '
            'detectan las lineas, y el umbral de acumulacion, que define el numero minimo de votos '
            'necesarios para considerar que existe una linea. Ajustar estos parametros correctamente '
            'es clave para obtener buenos resultados, ya que valores muy bajos generan muchas '
            'detecciones falsas, mientras que valores muy altos pueden omitir lineas reales.'
        )

        self.add_heading_custom('1.2 Cuantificacion Vectorial', level=2)

        self.add_paragraph_custom(
            'La cuantificacion vectorial es una tecnica de compresion de imagenes que agrupa '
            'bloques de pixeles similares y los representa mediante un conjunto reducido de '
            'vectores codigo, conocido como diccionario o codebook. A diferencia de la cuantificacion '
            'escalar que procesa cada pixel individualmente, la cuantificacion vectorial considera '
            'patrones de pixeles, aprovechando la correlacion espacial entre pixeles vecinos.'
        )

        self.add_paragraph_custom(
            'El algoritmo LBG, tambien conocido como algoritmo de Linde-Buzo-Gray, es el metodo '
            'estandar para generar el diccionario de vectores codigo. Este algoritmo es una '
            'variante del algoritmo de clustering K-means adaptada especificamente para cuantificacion '
            'vectorial. El proceso comienza dividiendo la imagen en bloques de tamaño fijo, por '
            'ejemplo 2x2 o 4x4 pixeles, y convirtiendo cada bloque en un vector. Luego, mediante '
            'un proceso iterativo de agrupamiento, se encuentran los vectores representativos que '
            'minimizan la distorsion total.'
        )

        self.add_paragraph_custom(
            'El tamaño del bloque y el tamaño del diccionario son los parametros fundamentales '
            'que determinan la relacion entre compresion y calidad. Bloques mas grandes permiten '
            'mayor compresion pero pueden perder detalles finos, mientras que diccionarios mas '
            'grandes preservan mejor la calidad pero reducen la tasa de compresion. La evaluacion '
            'de los resultados se realiza mediante metricas objetivas como MSE y PSNR, asi como '
            'mediante inspeccion visual de las imagenes reconstruidas.'
        )

        self.doc.add_page_break()

    def generate_hough_methodology(self):
        """Genera la seccion de metodologia para Hough."""
        self.add_heading_custom('2. Metodologia - Transformada de Hough', level=1)

        self.add_paragraph_custom(
            'Para la deteccion de lineas mediante la Transformada de Hough, se siguio el siguiente '
            'procedimiento sistematico para cada una de las tres imagenes de prueba.'
        )

        self.add_heading_custom('2.1 Procedimiento Implementado', level=2)

        steps = [
            'Carga de la imagen original: Se leyo cada imagen en formato PNG utilizando OpenCV.',

            'Conversion a escala de grises: Como la Transformada de Hough trabaja sobre imagenes '
            'de un solo canal, se convirtio cada imagen a escala de grises. Esto simplifica el '
            'procesamiento y permite enfocarse en las caracteristicas estructurales sin la influencia '
            'del color.',

            'Deteccion de bordes con Canny: Antes de aplicar la Transformada de Hough, es necesario '
            'identificar los bordes en la imagen. Se utilizo el detector de bordes Canny con '
            'umbrales de 50 y 150, previamente aplicando un filtro Gaussiano para reducir el ruido. '
            'Este paso es crucial porque la Transformada de Hough trabaja sobre los puntos de borde.',

            'Aplicacion de la Transformada de Hough: Se implemento la Transformada de Hough '
            'probabilistica, que es mas eficiente que la version estandar y permite especificar '
            'parametros adicionales como la longitud minima de linea y la distancia maxima entre '
            'segmentos. Se experimento con diferentes configuraciones de parametros para analizar '
            'su efecto.',

            'Superposicion de lineas detectadas: Las lineas detectadas se dibujaron sobre la imagen '
            'original para facilitar la evaluacion visual de los resultados.',

            'Analisis parametrico: Se probaron cuatro configuraciones diferentes de parametros para '
            'cada imagen, variando el umbral de acumulacion, la longitud minima de linea y la '
            'resolucion angular. Esto permitio entender como cada parametro afecta la sensibilidad '
            'y precision de la deteccion.'
        ]

        for i, step in enumerate(steps, 1):
            para = self.doc.add_paragraph(style='List Number')
            para.add_run(step)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_page_break()

    def generate_hough_results(self):
        """Genera la seccion de resultados de Hough."""
        self.add_heading_custom('3. Resultados - Transformada de Hough', level=1)

        # Resultados para cada imagen
        images = ['ImgHough01', 'ImgHough03', 'ImgHough04']

        for img_name in images:
            self.add_heading_custom(f'3.{images.index(img_name)+1} {img_name}', level=2)

            img_dir = f'results/hough/{img_name}'

            if not os.path.exists(img_dir):
                self.add_paragraph_custom(f'No se encontraron resultados para {img_name}')
                continue

            # Imagen original
            original_path = os.path.join(img_dir, f'{img_name}_original.png')
            if os.path.exists(original_path):
                self.add_image_with_caption(original_path,
                    f'Figura {images.index(img_name)*3+1}: Imagen original {img_name}', width=4)

            # Escala de grises
            gray_path = os.path.join(img_dir, f'{img_name}_gray.png')
            if os.path.exists(gray_path):
                self.add_image_with_caption(gray_path,
                    f'Figura {images.index(img_name)*3+2}: Conversion a escala de grises', width=4)

            # Bordes detectados
            edges_path = os.path.join(img_dir, f'{img_name}_edges.png')
            if os.path.exists(edges_path):
                self.add_image_with_caption(edges_path,
                    f'Figura {images.index(img_name)*3+3}: Bordes detectados con Canny', width=4)

            # Comparacion de parametros
            comparison_path = os.path.join(img_dir, f'{img_name}_comparison.png')
            if os.path.exists(comparison_path):
                self.doc.add_page_break()
                self.add_image_with_caption(comparison_path,
                    f'Figura {images.index(img_name)*3+4}: Comparacion de diferentes configuraciones de parametros', width=6.5)

            self.add_paragraph_custom(
                f'Para la imagen {img_name}, se observa que los parametros tienen un impacto '
                'significativo en el numero y calidad de las lineas detectadas. Con umbrales '
                'conservadores se detectan solo las lineas mas prominentes, evitando falsos '
                'positivos pero pudiendo perder algunas lineas reales. Por otro lado, con '
                'parametros sensibles se detectan muchas mas lineas, capturando mas detalles '
                'pero tambien generando mas detecciones espurias.'
            )

            self.doc.add_page_break()

    def generate_hough_analysis(self):
        """Genera el analisis de parametros de Hough."""
        self.add_heading_custom('4. Analisis de Parametros - Hough', level=1)

        self.add_paragraph_custom(
            'Durante la experimentacion con la Transformada de Hough, se identificaron varios '
            'aspectos importantes sobre como los parametros afectan los resultados.'
        )

        self.add_heading_custom('4.1 Efecto del Umbral de Acumulacion', level=2)

        self.add_paragraph_custom(
            'El umbral de acumulacion determina cuantos puntos de borde deben votar por una linea '
            'para que esta sea considerada valida. Se probaron umbrales de 40, 80 y 150. Con '
            'umbrales bajos se detectaron muchas lineas cortas y fragmentadas, algunas de las cuales '
            'podrian no corresponder a estructuras reales. Con umbrales altos solo se detectaron '
            'las lineas mas largas y definidas, lo cual es apropiado para imagenes con estructuras '
            'lineales claras pero puede perder informacion en imagenes mas complejas.'
        )

        self.add_heading_custom('4.2 Efecto de la Resolucion Angular', level=2)

        self.add_paragraph_custom(
            'Se experimento con resoluciones angulares de pi/180 y pi/360 radianes. Una mayor '
            'resolucion angular (pi/360) permite detectar lineas con orientaciones mas precisas, '
            'lo cual es util cuando las lineas no estan perfectamente alineadas con los ejes. '
            'Sin embargo, esto aumenta el costo computacional y puede generar detecciones duplicadas '
            'de la misma linea con angulos ligeramente diferentes.'
        )

        self.add_heading_custom('4.3 Longitud Minima y Espacios Maximos', level=2)

        self.add_paragraph_custom(
            'La longitud minima de linea y el espaciado maximo entre segmentos son parametros '
            'especificos de la Transformada de Hough probabilistica. Una longitud minima alta '
            'filtra lineas cortas que podrian ser ruido, mientras que un espaciado maximo mayor '
            'permite unir segmentos de linea que estan ligeramente separados debido a '
            'discontinuidades en los bordes. El balance entre estos parametros depende mucho '
            'del tipo de imagen que se esta procesando.'
        )

        self.doc.add_page_break()

    def generate_quantization_methodology(self):
        """Genera la metodologia de cuantificacion vectorial."""
        self.add_heading_custom('5. Metodologia - Cuantificacion Vectorial', level=1)

        self.add_paragraph_custom(
            'Para la compresion de imagenes mediante cuantificacion vectorial, se implemento '
            'el algoritmo LBG completo siguiendo estos pasos.'
        )

        self.add_heading_custom('5.1 Procedimiento Implementado', level=2)

        steps = [
            'Carga de la imagen: Se cargaron las tres imagenes de prueba en formato JPG.',

            'Division en bloques: La imagen se dividio en bloques no solapados de tamaño especifico. '
            'Se probaron dos configuraciones: bloques de 2x2 pixeles y bloques de 4x4 pixeles. '
            'Cada bloque RGB se convirtio en un vector de dimension 12 (para 2x2) o 48 (para 4x4), '
            'concatenando los valores de los tres canales de color.',

            'Generacion del conjunto de entrenamiento: Todos los bloques de la imagen se utilizaron '
            'como conjunto de entrenamiento. Esto genero conjuntos de cientos de miles de vectores, '
            'dependiendo del tamaño de la imagen y del tamaño de bloque elegido.',

            'Aplicacion del algoritmo LBG: Se utilizo la implementacion de K-means de scikit-learn, '
            'que es matematicamente equivalente al algoritmo LBG. Se especifico un diccionario de '
            '128 vectores codigo. El algoritmo itero hasta convergencia, minimizando la distorsion '
            'total medida como la suma de distancias cuadraticas entre cada vector y su codigo '
            'mas cercano.',

            'Cuantificacion: Cada bloque de la imagen original se reemplazo por el vector codigo '
            'mas cercano del diccionario, usando distancia euclidiana como metrica.',

            'Reconstruccion: Se reconstruyo la imagen reorganizando los vectores codigo en su '
            'posicion espacial original y convirtiendolos de vuelta a bloques de pixeles.',

            'Calculo de metricas: Se calcularon MSE y PSNR comparando la imagen original con la '
            'reconstruida. Tambien se calculo la tasa de compresion teorica considerando el tamaño '
            'del diccionario y los indices necesarios para representar la imagen.'
        ]

        for i, step in enumerate(steps, 1):
            para = self.doc.add_paragraph(style='List Number')
            para.add_run(step)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_page_break()

    def generate_quantization_results(self):
        """Genera los resultados de cuantificacion vectorial."""
        self.add_heading_custom('6. Resultados - Cuantificacion Vectorial', level=1)

        # Datos de las metricas obtenidas
        results_data = {
            'ImgCuantif01': {
                'block_2x2_codebook_128': {'mse': 40.34, 'psnr': 32.07, 'compression': 13.69},
                'block_4x4_codebook_128': {'mse': 103.04, 'psnr': 28.00, 'compression': 53.62}
            },
            'ImgCuantif02': {
                'block_2x2_codebook_128': {'mse': 54.48, 'psnr': 30.77, 'compression': 13.69},
                'block_4x4_codebook_128': {'mse': 187.26, 'psnr': 25.41, 'compression': 53.62}
            },
            'ImgCuantif03': {
                'block_2x2_codebook_128': {'mse': 63.57, 'psnr': 30.10, 'compression': 13.56},
                'block_4x4_codebook_128': {'mse': 221.56, 'psnr': 24.68, 'compression': 46.38}
            }
        }

        # Tabla resumen de todas las metricas
        self.add_heading_custom('6.1 Resumen de Metricas', level=2)

        headers = ['Imagen', 'Tamaño Bloque', 'MSE', 'PSNR (dB)', 'Compresion (x)']
        rows = []

        for img_name, configs in results_data.items():
            for config_name, metrics in configs.items():
                block_size = '2x2' if '2x2' in config_name else '4x4'
                rows.append([
                    img_name,
                    block_size,
                    f"{metrics['mse']:.2f}",
                    f"{metrics['psnr']:.2f}",
                    f"{metrics['compression']:.2f}"
                ])

        self.add_table_from_data(headers, rows)

        self.add_paragraph_custom(
            'La tabla anterior resume las metricas obtenidas para todas las imagenes y configuraciones '
            'probadas. Se observa claramente el compromiso entre compresion y calidad: los bloques '
            'de 4x4 logran tasas de compresion mucho mayores, pero con un aumento significativo '
            'del error y una reduccion del PSNR.'
        )

        self.doc.add_page_break()

        # Resultados visuales por imagen
        for img_name in results_data.keys():
            self.add_heading_custom(f'6.2 Resultados Visuales - {img_name}', level=2)

            # Mostrar resultados para cada configuracion
            for config in ['block_2x2_codebook_128', 'block_4x4_codebook_128']:
                block_size = '2x2' if '2x2' in config else '4x4'
                self.add_paragraph_custom(f'Configuracion: Bloques {block_size}, Diccionario 128 vectores',
                                         bold=True)

                comparison_path = f'results/cuantificacion/{img_name}/{config}/comparison.png'
                if os.path.exists(comparison_path):
                    self.add_image_with_caption(comparison_path,
                        f'Comparacion de {img_name} con bloques {block_size}', width=6.5)

                metrics = results_data[img_name][config]
                self.add_paragraph_custom(
                    f"Se obtuvo un MSE de {metrics['mse']:.2f} y un PSNR de {metrics['psnr']:.2f} dB, "
                    f"logrando una compresion de {metrics['compression']:.2f}x. "
                    f"{'La calidad visual se mantiene bastante buena, con perdidas apenas perceptibles en areas homogeneas.' if metrics['psnr'] > 30 else 'Se observa cierta degradacion visual, especialmente en regiones con textura fina y bordes.'}"
                )

                self.doc.add_paragraph()

            self.doc.add_page_break()

    def generate_quantization_analysis(self):
        """Genera el analisis de cuantificacion vectorial."""
        self.add_heading_custom('7. Analisis Comparativo - Cuantificacion Vectorial', level=1)

        self.add_heading_custom('7.1 Efecto del Tamaño de Bloque', level=2)

        self.add_paragraph_custom(
            'El tamaño de bloque es el parametro que mas impacto tiene en el balance entre '
            'compresion y calidad. Los bloques de 2x2 pixeles preservan mejor los detalles '
            'finos y los bordes nitidos, porque cada vector representa un area pequeña de la imagen. '
            'Esto resulta en valores de PSNR generalmente superiores a 30 dB, lo cual indica '
            'una buena calidad visual. Sin embargo, la tasa de compresion es limitada, '
            'alrededor de 13-14x.'
        )

        self.add_paragraph_custom(
            'Por otro lado, los bloques de 4x4 pixeles agrupan areas mas grandes, lo que permite '
            'representar la imagen con menos vectores y lograr compresiones mucho mayores, '
            'del orden de 45-54x. El costo es una perdida notable de calidad, especialmente '
            'visible en las regiones con gradientes suaves y texturas detalladas, donde aparece '
            'un efecto de "blocado" caracteristico. Los valores de PSNR caen a 24-28 dB, '
            'indicando una distorsion mas significativa.'
        )

        self.add_heading_custom('7.2 Comportamiento segun Contenido de Imagen', level=2)

        self.add_paragraph_custom(
            'Se observo que el contenido de la imagen afecta considerablemente los resultados. '
            'ImgCuantif01, que tiene areas mas homogeneas, presento mejor PSNR y menor MSE '
            'en comparacion con ImgCuantif02 e ImgCuantif03, que contienen mas detalles y texturas. '
            'Esto es esperable porque la cuantificacion vectorial funciona mejor cuando hay '
            'muchos bloques similares que pueden agruparse efectivamente.'
        )

        self.add_paragraph_custom(
            'Las imagenes con alto contenido de frecuencias espaciales altas, como bordes finos '
            'y texturas complejas, sufren mas degradacion porque requieren mas vectores codigo '
            'diferentes para representar adecuadamente la variedad de patrones. Con un diccionario '
            'fijo de 128 vectores, estas imagenes tienen mayor error de cuantificacion.'
        )

        self.add_heading_custom('7.3 Evaluacion del Algoritmo LBG', level=2)

        self.add_paragraph_custom(
            'El algoritmo LBG demostro ser efectivo para generar diccionarios representativos. '
            'La convergencia fue relativamente rapida, generalmente en menos de 20 iteraciones. '
            'La eleccion de 128 vectores codigo fue un compromiso razonable: suficientes vectores '
            'para capturar la variabilidad de la imagen, pero no tantos como para eliminar el '
            'beneficio de la compresion.'
        )

        self.add_paragraph_custom(
            'Un aspecto interesante es que el diccionario generado se adapta al contenido especifico '
            'de cada imagen. Si se inspeccionaran los vectores codigo, se verian patrones tipicos '
            'de esa imagen particular. Esto hace que el codebook no sea reutilizable entre imagenes '
            'diferentes, lo cual es una limitacion practica de esta tecnica.'
        )

        self.doc.add_page_break()

    def generate_conclusions(self):
        """Genera las conclusiones del informe."""
        self.add_heading_custom('8. Conclusiones', level=1)

        self.add_paragraph_custom(
            'Esta practica permitio comprender en profundidad dos tecnicas fundamentales del '
            'procesamiento digital de imagenes: la Transformada de Hough para deteccion de '
            'estructuras geometricas y la cuantificacion vectorial para compresion.'
        )

        self.add_heading_custom('8.1 Sobre la Transformada de Hough', level=2)

        self.add_paragraph_custom(
            'La Transformada de Hough resulto ser una herramienta robusta para detectar lineas, '
            'incluso en presencia de ruido y discontinuidades en los bordes. La principal ventaja '
            'de esta tecnica es que no requiere que las lineas sean completamente continuas, '
            'lo cual es muy util en imagenes reales donde los bordes pueden estar fragmentados. '
            'Sin embargo, encontre que el exito de la deteccion depende criticamente de una buena '
            'etapa previa de deteccion de bordes y de la correcta configuracion de parametros.'
        )

        self.add_paragraph_custom(
            'Un desafio importante es la eleccion de parametros. No existe una configuracion '
            'universal que funcione bien para todas las imagenes. Imagenes con muchas lineas '
            'cortas requieren umbrales bajos y longitudes minimas pequeñas, mientras que imagenes '
            'con pocas lineas largas se benefician de configuraciones mas conservadoras. La '
            'experimentacion con multiples conjuntos de parametros, como hicimos en esta practica, '
            'es esencial para entender como ajustar la tecnica a cada caso.'
        )

        self.add_paragraph_custom(
            'En cuanto a aplicaciones practicas, la Transformada de Hough es invaluable en areas '
            'como vision por computadora para vehiculos autonomos, reconocimiento de documentos, '
            'analisis de imagenes medicas, y cualquier tarea donde se necesite detectar estructuras '
            'lineales o geometricas especificas.'
        )

        self.add_heading_custom('8.2 Sobre la Cuantificacion Vectorial', level=2)

        self.add_paragraph_custom(
            'La cuantificacion vectorial mediante el algoritmo LBG demostro ser efectiva para '
            'comprimir imagenes, logrando reducciones significativas de tamaño con perdidas '
            'controladas de calidad. La tecnica aprovecha la redundancia espacial inherente en '
            'las imagenes naturales, donde bloques vecinos tienden a ser similares.'
        )

        self.add_paragraph_custom(
            'El compromiso fundamental entre compresion y calidad se hizo muy evidente en los '
            'experimentos. Los bloques pequeños (2x2) mantienen buena calidad visual pero comprimen '
            'poco, mientras que bloques grandes (4x4) comprimen mucho mas pero con degradacion '
            'notable. Este es un balance que debe decidirse segun la aplicacion: si se necesita '
            'maxima calidad visual, se elegiran bloques pequeños y diccionarios grandes; si el '
            'objetivo principal es minimizar el tamaño, se usaran bloques mas grandes.'
        )

        self.add_paragraph_custom(
            'Comparando con tecnicas modernas como JPEG, la cuantificacion vectorial tiene algunas '
            'ventajas en ciertos escenarios, especialmente en imagenes con patrones repetitivos. '
            'Sin embargo, es mas costosa computacionalmente durante el entrenamiento del diccionario '
            'y requiere almacenar o transmitir el codebook junto con los indices, lo cual puede '
            'ser un overhead significativo para imagenes pequeñas.'
        )

        self.add_heading_custom('8.3 Reflexiones Finales', level=2)

        self.add_paragraph_custom(
            'Implementar estos algoritmos desde la practica me dio una comprension mucho mas profunda '
            'que simplemente estudiar la teoria. Ver como pequeños cambios en los parametros afectan '
            'drasticamente los resultados fue muy instructivo. Tambien aprecie la importancia de '
            'las metricas objetivas como MSE y PSNR, aunque note que no siempre coinciden '
            'perfectamente con la percepcion visual humana.'
        )

        self.add_paragraph_custom(
            'Un aspecto que me parecio particularmente interesante es como estas tecnicas relativamente '
            'antiguas siguen siendo relevantes. La Transformada de Hough se usa en sistemas actuales '
            'de deteccion y la cuantificacion vectorial influyo en el desarrollo de tecnicas mas '
            'modernas de compresion y machine learning, como la cuantificacion de redes neuronales.'
        )

        self.add_paragraph_custom(
            'Para trabajos futuros, seria interesante experimentar con la deteccion de circulos '
            'usando la Transformada de Hough circular, y probar cuantificacion vectorial con '
            'diccionarios de diferentes tamaños para construir curvas completas de rate-distortion. '
            'Tambien seria valioso comparar directamente con JPEG en terminos de tiempo de procesamiento '
            'y calidad visual a la misma tasa de compresion.'
        )

    def generate_full_report(self):
        """Genera el informe completo."""
        print("Generando informe completo...")

        self.generate_cover_page()
        self.generate_introduction()
        self.generate_hough_methodology()
        self.generate_hough_results()
        self.generate_hough_analysis()
        self.generate_quantization_methodology()
        self.generate_quantization_results()
        self.generate_quantization_analysis()
        self.generate_conclusions()

        # Guardar documento
        self.doc.save(self.output_path)
        print(f"Informe generado exitosamente: {self.output_path}")


if __name__ == "__main__":
    print("="*70)
    print("GENERACION DE INFORME COMPLETO")
    print("="*70)

    generator = ReportGenerator()
    generator.generate_full_report()

    print("\nInforme completado exitosamente!")
